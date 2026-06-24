"""文档解析器 —— 基于 LangChain document_loaders，支持多格式

内置格式：txt / py / md / pdf / docx / xlsx 等。
parse_file() → (content: str | None, error: str | None)，保持向后兼容。
load_documents() → List[Document]，用于 LangChain 原生管线。
"""

import os
import chardet

from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document

# ================================================================
#  可识别为纯文本的扩展名
# ================================================================

_TEXT_EXTENSIONS: set[str] = {
    '.txt', '.py', '.json', '.xml', '.csv', '.md', '.yaml', '.yml',
    '.ini', '.cfg', '.toml', '.log', '.html', '.css', '.js', '.ts',
    '.java', '.c', '.cpp', '.h', '.rs', '.go', '.sh', '.bat', '.ps1',
    '.sql', '.r', '.rb', '.php', '.swift', '.kt', '.scala', '.lua',
}

# 单次解析最大字符数（避免撑爆上下文）
MAX_CHARS = 8000


# ================================================================
#  公共接口
# ================================================================

def parse_file(filepath: str) -> tuple[str | None, str | None]:
    """解析任意文件，返回 (内容, 错误)。

    内容可能被截断至 MAX_CHARS；错误为 None 表示成功。
    """
    if not os.path.isfile(filepath):
        return None, f"文件不存在: {filepath}"

    ext = os.path.splitext(filepath)[1].lower()

    try:
        if ext in _TEXT_EXTENSIONS:
            content = _read_text_file(filepath)
        elif ext == '.docx':
            content = _parse_docx(filepath)
        elif ext == '.pdf':
            content = _parse_pdf(filepath)
        elif ext in ('.xlsx', '.xls'):
            content = _parse_xlsx(filepath)
        elif ext in ('.doc',):
            return None, "暂不支持旧版 .doc 格式，请转换为 .docx 后重试"
        else:
            return None, f"暂不支持解析「{ext}」格式的文件"
    except Exception as e:
        return None, f"解析出错: {e}"

    if content and len(content) > MAX_CHARS:
        content = content[:MAX_CHARS] + f"\n\n...（内容已截断，共 {len(content)} 字符）"

    return content or None, None


def load_documents(filepath: str) -> list[Document]:
    """加载文件并返回 LangChain Document 对象列表。

    每个 Document 带 metadata（source、page 等），可直接喂给
    RecursiveCharacterTextSplitter.split_documents()。
    """
    if not os.path.isfile(filepath):
        raise FileNotFoundError(f"文件不存在: {filepath}")

    ext = os.path.splitext(filepath)[1].lower()

    if ext in _TEXT_EXTENSIONS:
        text = _read_text_file(filepath)
        if not text.strip():
            return []
        return [Document(
            page_content=text,
            metadata={"source": filepath, "format": ext.lstrip('.')},
        )]
    elif ext == '.pdf':
        docs = PyPDFLoader(filepath).load()
    elif ext == '.docx':
        docs = _DocxLoader(filepath).load()
    elif ext in ('.xlsx', '.xls'):
        docs = _ExcelLoader(filepath).load()
    else:
        raise ValueError(f"暂不支持解析「{ext}」格式的文件")

    for doc in docs:
        doc.page_content = doc.page_content.strip()
    return [d for d in docs if d.page_content]


# ================================================================
#  编码检测（chardet，自动识别 UTF-8 / GBK 等）
# ================================================================

def _read_text_file(filepath: str) -> str:
    """读取文本文件，自动检测编码（chardet）。

    对 Windows 中文环境（GBK）和 UTF-8 文件均能正确处理。
    """
    with open(filepath, 'rb') as f:
        raw = f.read()
    if not raw:
        return ""
    detected = chardet.detect(raw)
    encoding = detected.get('encoding') or 'utf-8'
    # chardet 对中文有时误判为 GB2312（是 GBK 的子集），统一用 GBK
    if encoding.upper() in ('GB2312', 'GB18030'):
        encoding = 'gbk'
    confidence = detected.get('confidence', 0)
    # 低置信度 → 用 UTF-8 兜底
    if confidence is not None and confidence < 0.6:
        encodings_to_try = [encoding, 'utf-8', 'gbk', 'latin-1']
    else:
        encodings_to_try = [encoding, 'utf-8', 'gbk', 'latin-1']
    for enc in encodings_to_try:
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    # 最终兜底
    return raw.decode('utf-8', errors='ignore')


# ================================================================
#  内部解析函数
# ================================================================

def _parse_docx(filepath: str) -> str:
    """Word 文档 (.docx) → python-docx"""
    from docx import Document as DocxDocument

    doc = DocxDocument(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return '\n'.join(paragraphs)


def _parse_pdf(filepath: str) -> str:
    """PDF → LangChain PyPDFLoader"""
    docs = PyPDFLoader(filepath).load()
    return '\n'.join(d.page_content for d in docs)


def _parse_xlsx(filepath: str) -> str:
    """Excel (.xlsx / .xls) → openpyxl"""
    from openpyxl import load_workbook

    wb = load_workbook(filepath, read_only=True)
    lines: list[str] = []

    for name in wb.sheetnames:
        ws = wb[name]
        lines.append(f"=== Sheet: {name} ===")
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            row_str = '\t'.join(str(c) if c is not None else '' for c in row)
            if row_str.strip():
                lines.append(row_str)
                row_count += 1
                if row_count > 200:
                    lines.append("...（表格行数过多，已截断）")
                    break
        if row_count == 0:
            lines.append("（空白工作表）")

    wb.close()
    return '\n'.join(lines)


# ================================================================
#  自定义 Loader —— docx / xlsx 返回统一 Document 对象
# ================================================================

class _DocxLoader:
    """轻量 docx → List[Document]，基于 python-docx（无需 docx2txt）"""

    def __init__(self, filepath: str):
        self._path = filepath

    def load(self) -> list[Document]:
        from docx import Document as DocxDocument

        doc = DocxDocument(self._path)
        text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        if not text:
            return []
        return [Document(
            page_content=text,
            metadata={"source": self._path, "format": "docx"},
        )]


class _ExcelLoader:
    """轻量 xlsx → List[Document]，基于 openpyxl（无需 unstructured）"""

    def __init__(self, filepath: str):
        self._path = filepath

    def load(self) -> list[Document]:
        from openpyxl import load_workbook

        wb = load_workbook(self._path, read_only=True)
        docs: list[Document] = []

        for name in wb.sheetnames:
            ws = wb[name]
            rows: list[str] = []
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                row_str = '\t'.join(str(c) if c is not None else '' for c in row)
                if row_str.strip():
                    rows.append(row_str)
                    row_count += 1
                    if row_count > 200:
                        rows.append("...（表格行数过多，已截断）")
                        break
            if rows:
                docs.append(Document(
                    page_content=f"=== Sheet: {name} ===\n" + '\n'.join(rows),
                    metadata={"source": self._path, "sheet": name, "format": "xlsx"},
                ))

        wb.close()
        return docs
